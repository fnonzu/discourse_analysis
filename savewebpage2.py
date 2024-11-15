import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import os
from docx import Document
from docx.shared import Inches
from PIL import Image  
import io

def scrape_article_to_word(url, save_folder, selectors):
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Extract title
    title = soup.select_one(selectors['title']).get_text() if soup.select_one(selectors['title']) else 'No Title'
    doc.add_heading(title, level=1)

    # Scrape content
    content = soup.select_one(selectors['content'])
    if content:
        paragraphs = content.find_all('p')
        for para in paragraphs:
            doc.add_paragraph(para.get_text())
    else:
        print("Content not found. Check selectors.")
        return

    # Download and add images
    img_tags = content.find_all('img') if content else []
    for img in img_tags:
        img_url = urljoin(url, img.get('src'))
        if img_url and img_url.startswith(('http:', 'https:')):
            img_data = requests.get(img_url).content
            img_stream = io.BytesIO(img_data)

            # Verify image format using PIL
            try:
                with Image.open(img_stream) as img_check:
                    img_format = img_check.format
                    if img_format in ["JPEG", "PNG"]:  # Supported formats
                        img_name = os.path.join(save_folder, os.path.basename(img_url))
                        with open(img_name, 'wb') as img_file:
                            img_file.write(img_data)
                        doc.add_picture(img_name, width=Inches(4))
                    else:
                        print(f"Unsupported image format for {img_url}. Skipping...")
            except Exception as e:
                print(f"Error processing image {img_url}: {e}")

    # Save the document
    filename = os.path.join(save_folder, f"{title}.docx")
    doc.save(filename)
    print(f"Saved: {filename}")

# Usage
url = "https://nikvesti.com/ua/articles/273906"
save_folder = r"YOUR PATH"
# selectors = {
#     "title": ["h1", ".article-title", ".headline", ".entry-title"],
#     "content": ["article", ".post-content", ".entry-content", ".content-body", ".text"],
#     "author": [".author", ".byline", ".author-name", ".writer"],
#     "date": ["time", ".date", "meta[property='article:published_time']", "meta[name='date']"],
#     "images": [".content img", ".article-body img", ".entry-content img"]
# }
selectors = {
    "title": ["h1"],
    "content": ["article"],
    "author": [".author"],
    "date": ["time"],
    "images": [".content img"]
}

scrape_article_to_word(url, save_folder, selectors)
